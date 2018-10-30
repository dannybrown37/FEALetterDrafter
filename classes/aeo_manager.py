import glob
import docx
from functions.validate import get_string, get_case_number


class AEOManager(object):
    def __init__(self, case_data=[]):
        # Get name of files and stripped allegation language from file names
        self.file_names = []
        self.allegations = []
        self.selections = []
        self.case_data = case_data
        if case_data:
            print "\nAnd now for the AEO..." # greeting of sorts from review_manager
            self.case_number = case_data[0]
            self.project_name = case_data[1]
        for file_name in glob.glob("AEOs/*.txt"): # relative to main, not class
            self.file_names.append(file_name)
            allegation = file_name.replace("AEOs\\", "").replace(".txt", "")
            self.allegations.append(allegation)
        # Get allegations user intends to use
        self.add_allegations()
        # Pull allegation language and stuff
        self.create_document()
        # TODO add dialogue box to select where to save? Figure out how to handle

    def add_allegations(self, search=""):
        if self.selections:
            print "\nYou have selected the following allegations so far:\n"
            for selection in self.selections:
                print str(selection).ljust(3), "|", self.allegations[selection]
            print "\nIf you're done selecting allegations, type \"done\"."
        if not search:
            prompt = "Enter a keyword to find the allegation."
            search = get_string(prompt).lower().strip()
        else:
            print "Searching for %s" % search
        if search == "done":
            return # breaks from recursive function
        possible_selections = []
        for i, (fn, al) in enumerate(zip(self.file_names, self.allegations)):
            with open(fn) as f:
                content = f.read()
            if search in al.lower() or search in content.lower():
                print str(i).ljust(3), "|", al
                possible_selections.append(i)
        if possible_selections:
            prompt = "Select a number or enter another search term."
            select = get_string(prompt)
        else:
            print "Not found!"
            self.add_allegations()
        try:
            select = int(select)
        except ValueError:
            search = select.lower()
            self.add_allegations(search=search)
        if select in possible_selections:
            self.selections.append(select)
            self.add_allegations()

    def create_document(self):
        if not self.case_data:
            self.case_number = get_case_number()
            self.project_name = get_string("Enter the respondent or project.")

        # Creating our .docx file
        last = self.selections[len(self.selections)-1]
        document = docx.Document()
        document.add_heading(
            "AEO: %s, %s" % (self.case_number, self.project_name), 0
        )
        doc_data = []
        for num in self.selections:
            with open(self.file_names[num], "r") as f:
                lines = f.readlines()
                f.close()
                bolden = [
                    "Allegation",
                    "Statute/Rule Reference",
                    "Elements",
                    "Evidence"
                ]
                p = document.add_paragraph("")
                p.add_run(lines[0]).bold = True
                for i in range(1, len(lines)):
                    if any(bold in lines[i] for bold in bolden):
                        p.add_run(lines[i].split(":")[0]).bold = True
                        p.add_run(":")
                        try:
                            p.add_run(lines[i].split(":")[1])
                        except:
                            print lines[i]
                    else:
                        p.add_run(lines[i])
                if num is not last:
                    p.add_run("\n\n_______________________________\n")
            doc_data.append(self.allegations[num].split(" ")[0])

        doc_data = " ".join(doc_data)

        document.save(" %s AEO %s.docx" % (doc_data, self.case_number))
